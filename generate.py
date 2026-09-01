import docx
from docx import Document
from docx.shared import Inches
from datetime import datetime

doc = Document()
doc.add_heading('Technical Documentation - Delivery Agent Chat System v2', 0)

doc.add_heading('Project Overview', level=1)
doc.add_paragraph('This is a real-time delivery agent chat and tracking platform built with Node.js, Express, Socket.IO, and React. It enables customers to track delivery agents on an interactive map, receive live ETA updates, engage in bidirectional chat, and get in-app notifications. Agents can update their location, which triggers real-time broadcasts and ETA recalculations using the Haversine formula (assuming 40 km/h average speed). The system uses in-memory storage for agents and customers and supports both REST and WebSocket communication.')

doc.add_heading('Users', level=1)
doc.add_paragraph('• Customers: View agent location on map, chat in real-time, receive ETA and notifications.\n• Agents: Update location, receive customer messages, see live ETA.')

doc.add_heading('Features', level=1)
doc.add_paragraph('• Real-time agent location tracking with map clicks for simulation.\n• Dynamic ETA calculation based on distance (Haversine) and speed model.\n• Bidirectional chat via Socket.IO with auto-notifications.\n• In-app notifications that auto-dismiss after 8 seconds.\n• Responsive Leaflet map with sidebar UI.\n• REST endpoints for health checks, location updates, and agent queries.\n• CORS enabled for cross-origin requests.')

doc.add_heading('Workflows', level=1)
doc.add_paragraph('1. Customer opens the app and joins via Socket.IO as "customer".\n2. Agent updates location via POST /api/agent/location or Socket "update-location" event → ETA is recalculated and broadcast to all clients.\n3. Customer clicks on the map to simulate agent movement or sends a chat message → updates propagate instantly (<2s).\n4. Notifications are emitted on new chats and auto-cleaned up.\n5. Graceful shutdown on SIGTERM.')

doc.add_heading('Data Model', level=1)
doc.add_paragraph('Agent object (stored in Map):\n- agentId (string)\n- location: {latitude, longitude}\n- status (e.g. "enroute")\n- eta (minutes, integer)\n- destination (optional coords)\n- updated (Date)\n\nMessage payload:\n- from, to, message, timestamp, agentId')

doc.add_heading('Business Rules', level=1)
doc.add_paragraph('• ETA = distance (km) / 40 km/h * 60, rounded to nearest minute.\n• Only valid location updates with agentId, lat, lng are accepted (400 error otherwise).\n• Chat messages are broadcast to the target role and trigger notifications for the agent.\n• In-memory state is not persisted across restarts.')

doc.add_heading('Constraints', level=1)
doc.add_paragraph('• Node.js >=18.\n• No external API keys (uses public OpenStreetMap tiles).\n• In-memory only (no database).\n• All clients share the same simulated agent for demo purposes.')

doc.add_heading('Non-Functional Requirements', level=1)
doc.add_paragraph('• Low-latency real-time updates (<500ms target, <2s observed).\n• Responsive UI that works on desktop and mobile.\n• Graceful shutdown handling.\n• CORS permissive for development (*).\n• No authentication in this v2 demo.')

doc.add_heading('Edge Cases', level=1)
doc.add_paragraph('• Missing location data → 400 error.\n• Unknown agent query → 404.\n• Multiple simultaneous location updates (last-write wins in Map).\n• Client disconnect cleanup for customers.\n• No destination provided → default ETA handling.\n• Rapid map clicks → frequent broadcasts (no rate limiting in current impl).')

doc.add_heading('Acceptance Criteria', level=1)
doc.add_paragraph('• Backend serves on :5000 with /health and location endpoints (verified).\n• Frontend builds/runs with interactive map, chat, notifications, and ETA display.\n• Location updates propagate in <2s and update ETA.\n• Chat works bidirectionally with notifications.\n• Map click simulates agent movement.\n• npm run install-all, npm start, and npm run client succeed.')

doc.add_heading('Non-Goals', level=1)
doc.add_paragraph('• Persistent storage or database integration.\n• Real multi-agent support or authentication/authorization.\n• Production rate limiting, error tracking, or scaling.\n• Mobile native app (web only).')

doc.add_heading('Architecture', level=1)
doc.add_heading('Modules / Layers', level=2)
doc.add_paragraph('• Backend (server.js): Express REST + Socket.IO server, in-memory state, ETA logic.\n• Frontend (React + react-leaflet): Map view, chat sidebar, Socket client.\n• No separate service layers; monolithic for v2 demo.')

doc.add_heading('Allowed / Forbidden Dependencies', level=2)
doc.add_paragraph('Allowed: express, socket.io, cors, dotenv, haversine-distance, react, react-leaflet, socket.io-client.\nForbidden in current scope: databases (e.g. mongoose), auth libraries, heavy state management (Redux).')

doc.add_heading('Data Flow', level=2)
doc.add_paragraph('1. Client → Socket "update-location" or POST /api/agent/location → server updates Map and broadcasts "agent-update".\n2. Client map click → emits update → UI + Socket sync.\n3. Chat message → server broadcasts to target and emits notification.\n4. Health checks via simple GET.')

doc.add_heading('Public Interfaces', level=2)
doc.add_paragraph('• REST: GET /api/health, POST /api/agent/location, GET /api/agent/:id.\n• Socket events: join, chat, update-location, agent-update, notification, chat-sent.')

doc.add_heading('Integration Points', level=2)
doc.add_paragraph('• Socket.IO for real-time (no external broker).\n• Leaflet + OpenStreetMap tiles (public).\n• In-memory only; easy to swap for Redis later.')

doc.add_heading('Files in Scope', level=1)
doc.add_paragraph('• README.md (living spec)\n• package.json / client/package.json\n• server.js (core backend)\n• client/public/index.html, client/src/App.js, client/src/index.js\n• spec.md (additional details)')

doc.add_heading('Setup & Run Instructions', level=1)
doc.add_paragraph('1. npm install\n2. npm run install-all\n3. npm start (backend on 5000)\n4. In separate terminal: npm run client (React on 3000)\n\nAccess at http://localhost:3000. Click map to move agent; type in chat.')

doc.add_heading('Version & Generation Info', level=1)
doc.add_paragraph(f'Generated on: {datetime.now().isoformat()}\nVersion: 2.0.0 (updated from previous hello-world and basic-chat docs)\nThis document is living and should be regenerated when behavior changes.')

doc.save('delivery-agent-chat-system-v2-technical-documentation.docx')
print('Document generated successfully: delivery-agent-chat-system-v2-technical-documentation.docx')